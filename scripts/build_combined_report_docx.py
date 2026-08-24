#!/usr/bin/env python3
"""Build a single Word document from the two Markdown course reports."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
COURSE_MD = ROOT / "reports" / "课程设计报告_Adaptive_FAST_LIO2.md"
SURVEY_MD = ROOT / "reports" / "行业调研报告_机器人_AI_芯片与大模型.md"
OUTPUT_DOCX = ROOT / "reports" / "智能系统与机器人学课程报告_Adaptive_FAST_LIO2_含调研报告.docx"


def set_run_font(run, east_asia: str = "宋体", western: str = "Times New Roman",
                 size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = western
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_run_font(run, size=9)


def add_inline_runs(paragraph, text: str) -> None:
    """Render a compact subset of Markdown inline formatting."""
    token_re = re.compile(r"(\*\*.+?\*\*|`.+?`|\*.+?\*|https?://\S+)")
    position = 0
    for match in token_re.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, size=10.5)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=10.5, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, east_asia="等线", western="Consolas", size=9.5)
            run.font.color.rgb = RGBColor(70, 70, 70)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=10.5)
            run.italic = True
        else:
            run = paragraph.add_run(token.rstrip(".,;，。；"))
            set_run_font(run, size=9.5)
            run.font.color.rgb = RGBColor(5, 99, 193)
            run.underline = True
            trailing = token[len(run.text):]
            if trailing:
                tail = paragraph.add_run(trailing)
                set_run_font(tail, size=10.5)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=10.5)


def add_body_paragraph(document: Document, text: str, *,
                       first_line: bool = True,
                       alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph = document.add_paragraph()
    paragraph.alignment = alignment
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(3)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    add_inline_runs(paragraph, text)
    return paragraph


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row_index, values in enumerate(rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = values[column_index] if column_index < len(values) else ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline_runs(paragraph, text)
            for run in paragraph.runs:
                set_run_font(run, size=8.5, bold=(row_index == 0))
            if row_index == 0:
                set_cell_shading(cell, "D9EAF7")
        if row_index == 0:
            set_repeat_table_header(table.rows[row_index])
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_image(document: Document, alt: str, source: str, base_dir: Path) -> None:
    image_path = (base_dir / source).resolve()
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if image_path.exists():
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(6.3))
        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(6)
        caption_run = caption.add_run(f"图　{alt}")
        set_run_font(caption_run, size=9)
    else:
        run = paragraph.add_run(f"[图片缺失：{alt}，{image_path}]")
        set_run_font(run, size=9)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, east_asia="黑体", size=15, bold=True)
    elif level == 2:
        set_run_font(run, east_asia="黑体", size=13, bold=True)
    else:
        set_run_font(run, east_asia="黑体", size=11, bold=True)


def add_title(document: Document, title: str, survey: bool) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(30 if survey else 45)
    paragraph.paragraph_format.space_after = Pt(24)
    run = paragraph.add_run(title)
    set_run_font(run, east_asia="黑体", size=20, bold=True)


def render_markdown(document: Document, markdown_path: Path, *, survey: bool = False) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []
    in_formula = False
    formula_lines: list[str] = []
    first_h1 = True

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if stripped.startswith("```"):
            if in_code:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.left_indent = Cm(0.6)
                paragraph.paragraph_format.right_indent = Cm(0.6)
                paragraph.paragraph_format.space_after = Pt(6)
                run = paragraph.add_run("\n".join(code_lines))
                set_run_font(run, east_asia="等线", western="Consolas", size=8.5)
                code_lines.clear()
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue

        if stripped == "$$":
            if in_formula:
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(3)
                run = paragraph.add_run(" ".join(line.strip() for line in formula_lines))
                set_run_font(run, east_asia="等线", western="Cambria Math", size=10.5)
                formula_lines.clear()
                in_formula = False
            else:
                in_formula = True
            index += 1
            continue
        if in_formula:
            formula_lines.append(raw)
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        image_match = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            add_image(document, image_match.group(1), image_match.group(2), markdown_path.parent)
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines):
            separator = lines[index + 1].strip()
            if re.fullmatch(r"\|?[\s:|-]+\|?", separator):
                rows = [parse_table_row(raw)]
                index += 2
                while index < len(lines) and lines[index].strip().startswith("|"):
                    rows.append(parse_table_row(lines[index]))
                    index += 1
                add_table(document, rows)
                continue

        if stripped.startswith("# "):
            title = stripped[2:].strip()
            if first_h1:
                add_title(document, title, survey)
                first_h1 = False
            else:
                add_heading(document, title, 1)
            index += 1
            continue
        if stripped.startswith("## "):
            add_heading(document, stripped[3:].strip(), 1)
            index += 1
            continue
        if stripped.startswith("### "):
            add_heading(document, stripped[4:].strip(), 2)
            index += 1
            continue

        if stripped.startswith("> "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(4)
            add_inline_runs(paragraph, stripped[2:])
            index += 1
            continue

        ordered_match = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if ordered_match:
            paragraph = document.add_paragraph(style="List Number")
            paragraph.paragraph_format.line_spacing = 1.5
            paragraph.paragraph_format.space_after = Pt(2)
            add_inline_runs(paragraph, ordered_match.group(2))
            index += 1
            continue

        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.line_spacing = 1.5
            paragraph.paragraph_format.space_after = Pt(2)
            add_inline_runs(paragraph, stripped[2:])
            index += 1
            continue

        first_line = not (
            stripped.startswith("**关键词")
            or re.match(r"^\[\d+\]", stripped)
        )
        add_body_paragraph(document, stripped, first_line=first_line)
        index += 1


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.4)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)

    for level in range(1, 4):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.color.rgb = RGBColor(0, 0, 0)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    settings = document.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def main() -> None:
    document = Document()
    configure_document(document)
    render_markdown(document, COURSE_MD)

    # The survey must begin on the page immediately following the course
    # report references, as required by the user.
    document.add_page_break()
    render_markdown(document, SURVEY_MD, survey=True)

    properties = document.core_properties
    properties.title = "智能系统与机器人学课程报告：Adaptive FAST-LIO2（含行业调研报告）"
    properties.subject = "课程设计报告与机器人、AI、芯片、大模型行业调研"
    properties.author = "__________"
    properties.keywords = "FAST-LIO2, SLAM, 机器人, 人工智能, AI芯片, 大模型"

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
